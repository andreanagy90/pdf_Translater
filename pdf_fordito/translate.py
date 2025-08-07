import deepl
from docx import Document
import convert, time


def translater ():
    api_key = "d6efc743-00d3-4883-a711-6772c66c8caf:fx"

    translater = deepl.Translator(api_key)

    doc = Document(convert.doc_file)
    new_doc = Document()
    total_pages = len(doc.paragraphs)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"Start translating page {i+1}/{total_pages}")
            result = translater.translate_text(text, source_lang="EN", target_lang="HU")
            new_doc.add_paragraph(result.text)
            print(f"Finished translating page {i+1}/{total_pages}")
            time.sleep(0.5)
        else:
            new_doc.add_paragraph("")

    output_file = convert.doc_file.replace(".docx", "_hu.docx")
    new_doc.save(output_file)

    print(f"Translated file saved {output_file}")

    return output_file